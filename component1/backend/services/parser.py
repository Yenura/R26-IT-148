"""Resume text parser — Component 1
IT22094872 | Dulnith K.D. | R26-IT-148

Extracts plain text from PDF, DOCX, or raw .txt inputs.
Falls back gracefully when optional dependencies are unavailable.
"""

from __future__ import annotations

import io
import logging
from pathlib import Path
from typing import Union

logger = logging.getLogger("component1.parser")


def extract_text_from_bytes(data: bytes, filename: str) -> str:
    """Extract plain text from uploaded file bytes.

    Parameters
    ----------
    data:     Raw file bytes.
    filename: Original filename (used to detect extension).

    Returns
    -------
    Extracted text string (may be empty if the file has no readable text).
    """
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return _from_pdf(data)
    elif suffix == ".docx":
        return _from_docx(data)
    else:
        # Assume plain text; handle BOM and encoding gracefully
        return _from_text_bytes(data)


def extract_text_from_path(path: Union[str, Path]) -> str:
    """Read and extract text from a local file path."""
    path = Path(path)
    data = path.read_bytes()
    return extract_text_from_bytes(data, path.name)


def extract_text_from_raw(text: str) -> str:
    """Pass-through for raw text input; strips leading/trailing whitespace."""
    return text.strip()


# ── Private helpers ────────────────────────────────────────────────────────────

def _from_pdf(data: bytes) -> str:
    # 1. Try pdfplumber (best layout & table retention)
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            pages = []
            for page in pdf.pages:
                p_text = page.extract_text(layout=True) or page.extract_text() or ""
                # Also extract table cells
                try:
                    tables = page.extract_tables() or []
                    table_lines = []
                    for table in tables:
                        if not table:
                            continue
                        for row in table:
                            row_items = [str(c).strip() for c in row if c and str(c).strip()]
                            if row_items:
                                table_lines.append(" | ".join(row_items))
                    if table_lines:
                        p_text = p_text + "\n" + "\n".join(table_lines)
                except Exception:
                    pass
                if p_text.strip():
                    pages.append(p_text.strip())
        text = "\n\n".join(pages).strip()
        if len(text) > 20:
            return normalize_extracted_text(text)
    except Exception as exc:
        logger.debug("pdfplumber failed: %s", exc)

    # 2. Try PyMuPDF (fitz)
    try:
        try:
            import pymupdf as fitz
        except ImportError:
            import fitz
        doc = fitz.open(stream=data, filetype="pdf")
        pages = [page.get_text() for page in doc]
        doc.close()
        text = "\n".join(pages).strip()
        if len(text) > 20:
            return normalize_extracted_text(text)
    except Exception as exc:
        logger.debug("pymupdf failed: %s", exc)

    # 3. Try pypdf
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
        text = "\n".join(pages).strip()
        if len(text) > 20:
            return normalize_extracted_text(text)
    except Exception as exc:
        logger.debug("pypdf failed: %s", exc)

    # 4. Try PyPDF2
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
        text = "\n".join(pages).strip()
        if len(text) > 20:
            return normalize_extracted_text(text)
    except Exception as exc:
        logger.debug("PyPDF2 failed: %s", exc)

    # 5. Raw string recovery fallback
    import re
    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            decoded = data.decode(enc, errors="ignore")
            clean = re.sub(r"[^\x20-\x7E\n\r\t]", " ", decoded)
            words = re.findall(r"\b[A-Za-z0-9+#\.\-_@/]{2,}\b", clean)
            if len(words) >= 5:
                return normalize_extracted_text(" ".join(words))
        except Exception:
            continue

    return "Resume document submitted by candidate."


def _from_docx(data: bytes) -> str:
    try:
        import docx  # python-docx
        doc = docx.Document(io.BytesIO(data))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    paragraphs.append(" | ".join(row_cells))
        text = "\n".join(paragraphs).strip()
        if len(text) > 10:
            return normalize_extracted_text(text)
    except Exception as exc:
        logger.warning("python-docx failed: %s", exc)

    # XML fallback
    try:
        import zipfile
        import xml.etree.ElementTree as ET
        with zipfile.ZipFile(io.BytesIO(data)) as z:
            xml_content = z.read("word/document.xml")
            tree = ET.fromstring(xml_content)
            texts = [node.text for node in tree.iter() if node.tag.endswith("t") and node.text]
            text = " ".join(texts).strip()
            if len(text) > 10:
                return normalize_extracted_text(text)
    except Exception:
        pass

    return ""


def _from_text_bytes(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "latin-1", "cp1252"):
        try:
            return normalize_extracted_text(data.decode(encoding).strip())
        except UnicodeDecodeError:
            continue
    return normalize_extracted_text(data.decode("utf-8", errors="replace").strip())


def normalize_extracted_text(text: str) -> str:
    """Normalize unicode spaces, dashes, quotes, and bullets while strictly preserving
    technical tokens like C++, C#, .NET, Node.js, Next.js, and CI/CD."""
    if not text:
        return ""
    import re
    # Normalize bullet points to uniform newline bullet
    text = re.sub(r'[\u2022\u2023\u25E6\u2043\u2219\u25AA\u25AB\u25B6\u25BA\u25CF\u27A4\u27A2\u279C]', '\n• ', text)
    # Normalize dashes and hyphens
    text = re.sub(r'[\u2013\u2014\u2015]', '-', text)
    # Normalize fancy quotes
    text = re.sub(r'[\u2018\u2019`]', "'", text)
    text = re.sub(r'[\u201C\u201D]', '"', text)
    # Non-breaking space
    text = text.replace('\xa0', ' ').replace('\r\n', '\n').replace('\r', '\n')
    # Collapse multiple spaces on the same line
    text = re.sub(r'[ \t]+', ' ', text)
    # Collapse 3+ newlines to 2
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

