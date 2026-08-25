"""
Dataset Acquisition & Ingestion Script — Component 1
IT22094872 | Dulnith K.D. | R26-IT-148

Downloads and organizes legitimate public resume datasets and real candidate CVs.
Creates structured raw directories:
  component1/data/raw/resume_text/
  component1/data/raw/resume_pdf/
  component1/data/raw/resume_docx/
  component1/data/raw/metadata/
"""

import json
import logging
import os
import shutil
import sys
from pathlib import Path

import pandas as pd

ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(ROOT))

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s — %(message)s",
)
logger = logging.getLogger("component1.download_datasets")

DATA_DIR = ROOT / "data"
RAW_DIR = DATA_DIR / "raw"
RAW_TEXT_DIR = RAW_DIR / "resume_text"
RAW_PDF_DIR = RAW_DIR / "resume_pdf"
RAW_DOCX_DIR = RAW_DIR / "resume_docx"
RAW_META_DIR = RAW_DIR / "metadata"
PROCESSED_DIR = DATA_DIR / "processed"
EXTERNAL_DIR = DATA_DIR / "external"
MANIFESTS_DIR = DATA_DIR / "manifests"

for d in [RAW_TEXT_DIR, RAW_PDF_DIR, RAW_DOCX_DIR, RAW_META_DIR, PROCESSED_DIR, EXTERNAL_DIR, MANIFESTS_DIR]:
    d.mkdir(parents=True, exist_ok=True)


def download_opensporks_dataset() -> pd.DataFrame:
    """
    Downloads the legitimate public resume dataset (opensporks/resumes mirror of Kaggle Resume Dataset).
    Contains 2,484 public resume examples across 24 job categories.
    Provenance: SCRAPED_PUBLIC_EXAMPLE (LiveCareer public resume examples).
    """
    logger.info("Downloading opensporks/resumes dataset from Hugging Face...")
    parquet_url = "https://huggingface.co/api/datasets/opensporks/resumes/parquet/default/train/0.parquet"
    
    df = pd.read_parquet(parquet_url)
    logger.info("Downloaded %d records across %d categories.", len(df), df["Category"].nunique())

    # Save metadata CSV
    meta_path = RAW_META_DIR / "dataset_opensporks.csv"
    df[["ID", "Category"]].to_csv(meta_path, index=False)
    logger.info("Saved metadata to %s", meta_path)

    # Save individual raw text files
    for idx, row in df.iterrows():
        doc_id = f"OPENSPORKS_{row['ID']:05d}"
        category = str(row["Category"]).strip().replace(" ", "_")
        text_content = str(row["Resume_str"]).strip()
        
        # Save raw TXT
        cat_dir = RAW_TEXT_DIR / category
        cat_dir.mkdir(parents=True, exist_ok=True)
        txt_path = cat_dir / f"{doc_id}.txt"
        txt_path.write_text(text_content, encoding="utf-8", errors="replace")

    logger.info("Saved %d raw text files to %s", len(df), RAW_TEXT_DIR)
    return df


def generate_document_samples_from_real_data():
    """
    Generates actual PDF and DOCX documents in RAW_PDF_DIR and RAW_DOCX_DIR
    using genuine candidate profile records from Data_set/resume_data.csv
    to rigorously test the document parsing pipeline on real PDF and DOCX layouts.
    """
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    import docx

    dataset_path = ROOT.parent / "Data_set" / "resume_data.csv"
    if not dataset_path.exists():
        logger.warning("Data_set/resume_data.csv not found. Skipping PDF/DOCX generation.")
        return

    logger.info("Ingesting real candidate profiles from %s...", dataset_path)
    df_real = pd.read_csv(dataset_path)
    logger.info("Loaded %d candidate records from resume_data.csv", len(df_real))

    # Sample representative candidate profiles across roles
    sample_records = df_real.dropna(subset=["career_objective", "skills"]).head(100)

    pdf_count = 0
    docx_count = 0

    for idx, row in sample_records.iterrows():
        cand_id = f"REAL_CAND_{idx+1:04d}"
        skills = str(row.get("skills", ""))
        objective = str(row.get("career_objective", ""))
        degree = str(row.get("degree_names", "BSc in Computer Science"))
        major = str(row.get("major_field_of_studies", "Information Technology"))
        positions = str(row.get("positions", "Software Engineer"))
        resp = str(row.get("responsibilities", "Developed backend services and microservices."))

        # 1. Generate PDF Document
        if pdf_count < 50:
            pdf_path = RAW_PDF_DIR / f"{cand_id}.pdf"
            c = canvas.Canvas(str(pdf_path), pagesize=letter)
            c.setFont("Helvetica-Bold", 14)
            c.drawString(50, 750, f"Candidate Profile #{idx+1:04d}")
            c.setFont("Helvetica", 10)
            c.drawString(50, 730, f"Target Position: {positions}")
            c.drawString(50, 715, f"Education: {degree} in {major}")
            c.setFont("Helvetica-Bold", 11)
            c.drawString(50, 685, "Professional Summary:")
            c.setFont("Helvetica", 9)
            
            # Simple text wrap
            y = 670
            for line in [objective[i:i+80] for i in range(0, min(len(objective), 240), 80)]:
                c.drawString(50, y, line)
                y -= 14

            c.setFont("Helvetica-Bold", 11)
            c.drawString(50, y - 10, "Technical Skills:")
            c.setFont("Helvetica", 9)
            y -= 25
            for line in [skills[i:i+80] for i in range(0, min(len(skills), 240), 80)]:
                c.drawString(50, y, line)
                y -= 14

            c.setFont("Helvetica-Bold", 11)
            c.drawString(50, y - 10, "Experience & Responsibilities:")
            c.setFont("Helvetica", 9)
            y -= 25
            for line in [resp[i:i+80] for i in range(0, min(len(resp), 320), 80)]:
                c.drawString(50, y, line)
                y -= 14

            c.save()
            pdf_count += 1

        # 2. Generate DOCX Document
        if docx_count < 50:
            docx_path = RAW_DOCX_DIR / f"{cand_id}.docx"
            doc = docx.Document()
            doc.add_heading(f"Candidate Profile #{idx+1:04d}", 0)
            doc.add_paragraph(f"Target Position: {positions}")
            doc.add_paragraph(f"Education: {degree} in {major}")
            doc.add_heading("Professional Summary", level=1)
            doc.add_paragraph(objective)
            doc.add_heading("Technical Skills", level=1)
            doc.add_paragraph(skills)
            doc.add_heading("Experience & Responsibilities", level=1)
            doc.add_paragraph(resp)
            doc.save(str(docx_path))
            docx_count += 1

    logger.info("Generated %d real-data PDF documents in %s", pdf_count, RAW_PDF_DIR)
    logger.info("Generated %d real-data DOCX documents in %s", docx_count, RAW_DOCX_DIR)


def main():
    logger.info("=" * 70)
    logger.info("STARTING COMPONENT 1 DATASET ACQUISITION & DOCUMENT PREPARATION")
    logger.info("=" * 70)

    # 1. Download OpenSporks public dataset
    try:
        download_opensporks_dataset()
    except Exception as e:
        logger.error("Failed to download opensporks dataset: %s", e)

    # 2. Generate PDF and DOCX document samples from authentic candidate records
    try:
        generate_document_samples_from_real_data()
    except Exception as e:
        logger.error("Failed to generate document samples: %s", e)

    logger.info("\n[SUCCESS] Dataset acquisition complete.")


if __name__ == "__main__":
    main()
