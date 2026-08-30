"""Resume file parsing and NLP entity extraction."""
import io
import os
import re
import sys
from datetime import datetime
from typing import Any

# Import Component 1 high-precision skill extraction engine
try:
    _cur_dir = os.path.dirname(os.path.abspath(__file__))
    _root_dir = os.path.dirname(os.path.dirname(os.path.dirname(_cur_dir)))
    _c1_path = os.path.join(_root_dir, "component1")
    if os.path.exists(_c1_path) and _c1_path not in sys.path:
        sys.path.insert(0, _c1_path)
    from ml.extractor import extract_skills_and_certifications as c1_extract_skills
except Exception as _e:
    c1_extract_skills = None


def parse_resume_file(content: bytes, filename: str) -> str:
    """Extract text from PDF, DOCX, or TXT files."""
    name_lower = filename.lower()
    text = ""
    if name_lower.endswith(".pdf"):
        text = _parse_pdf(content)
    elif name_lower.endswith((".docx", ".doc")):
        text = _parse_docx(content)
    
    if not text.strip():
        # Fallback to UTF-8 / latin-1 decoding
        for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252", "ascii"):
            try:
                decoded = content.decode(enc, errors="ignore")
                # Keep printable characters and clean newlines
                clean = re.sub(r"[^\x20-\x7E\n\r\t]", " ", decoded)
                lines = [l.strip() for l in clean.splitlines() if l.strip()]
                if len(lines) >= 1:
                    text = "\n".join(lines)
                    break
            except Exception:
                continue

    if not text.strip():
        text = f"Resume file {filename} submitted by candidate."

    return text.strip()


def _parse_pdf(content: bytes) -> str:
    # 1. Try pdfplumber (best layout & multi-column retention)
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
        text = "\n".join(pages).strip()
        if len(text) > 20:
            return text
    except Exception:
        pass

    # 2. Try pymupdf (fitz)
    try:
        try:
            import pymupdf as fitz
        except ImportError:
            import fitz
        doc = fitz.open(stream=content, filetype="pdf")
        pages = []
        for page in doc:
            pages.append(page.get_text())
        doc.close()
        text = "\n".join(pages).strip()
        if len(text) > 20:
            return text
    except Exception:
        pass

    # 3. Try pypdf / PyPDF2
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(content))
        pages = [p.extract_text() or "" for p in reader.pages]
        text = "\n".join(pages).strip()
        if len(text) > 20:
            return text
    except Exception:
        pass

    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(content))
        pages = [p.extract_text() or "" for p in reader.pages]
        text = "\n".join(pages).strip()
        if len(text) > 20:
            return text
    except Exception:
        pass

    # 4. Try pdfminer
    try:
        from pdfminer.high_level import extract_text
        text = extract_text(io.BytesIO(content)).strip()
        if len(text) > 20:
            return text
    except Exception:
        pass

    return ""


def _parse_docx(content: bytes) -> str:
    # 1. Try python-docx
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        paras = [p.text for p in doc.paragraphs if p.text.strip()]
        # also read tables in docx
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paras.append(cell.text.strip())
        text = "\n".join(paras).strip()
        if len(text) > 10:
            return text
    except Exception:
        pass

    # 2. Try raw XML extraction from docx zip archive
    try:
        import zipfile
        import xml.etree.ElementTree as ET
        with zipfile.ZipFile(io.BytesIO(content)) as z:
            xml_content = z.read("word/document.xml")
            tree = ET.fromstring(xml_content)
            # Find all text tags <w:t>
            texts = [node.text for node in tree.iter() if node.tag.endswith("t") and node.text]
            text = " ".join(texts).strip()
            if len(text) > 10:
                return text
    except Exception:
        pass

    return ""


# ── NLP Preprocessing ──────────────────────────────────────────
STOPWORDS = frozenset("a an the and or but in on at to for of with is it its this that was were be been "
                      "being have has had do does did will would could should may might can shall from "
                      "by as into through during before after above below between out off over under "
                      "again further then once here there when where why how all each every both few "
                      "more most other some such no nor not only own same so than too very s t just "
                      "don now d ll m o re ve y ain aren couldn didn doesn hadn hasn haven isn ma "
                      "mightn mustn needn shan shouldn wasn weren won wouldn".split())

_CLEAN_RE = re.compile(r"[^a-z0-9\s#+./]")
_SPACE_RE = re.compile(r"\s+")
_lemmatizer = None

def _get_lemmatizer():
    global _lemmatizer
    if _lemmatizer is None:
        try:
            from nltk.stem import WordNetLemmatizer
            _lemmatizer = WordNetLemmatizer()
        except (ImportError, LookupError):
            _lemmatizer = False
    return _lemmatizer if _lemmatizer is not False else None


def preprocess_text(text: str) -> str:
    """Lowercase, remove punctuation, remove stopwords, tokenize, lemmatize."""
    if not text:
        return ""
    text_clean = _CLEAN_RE.sub(" ", text.lower())
    text_clean = _SPACE_RE.sub(" ", text_clean).strip()
    tokens = text_clean.split()
    tokens = [t for t in tokens if t not in STOPWORDS and len(t) > 1]
    lemmatizer = _get_lemmatizer()
    if lemmatizer:
        try:
            tokens = [lemmatizer.lemmatize(t) for t in tokens]
        except Exception:
            pass
    return " ".join(tokens)


# ── Entity Extraction ───────────────────────────────────────────
SKILLS_KEYWORDS = [
    # Languages
    "python", "java", "javascript", "typescript", "c++", "c#", ".net", "golang", "go", "rust", "kotlin", "swift", "php", "ruby", "r", "scala", "dart", "c",
    # Frontend
    "react", "react.js", "reactjs", "angular", "angularjs", "vue", "vue.js", "vuejs", "next.js", "nextjs", "nuxt.js", "svelte",
    "html", "html5", "css", "css3", "sass", "less", "tailwind", "tailwind css", "tailwindcss", "bootstrap", "redux", "webpack", "vite", "responsive design",
    # Backend & APIs
    "node.js", "nodejs", "express", "express.js", "expressjs", "django", "flask", "fastapi", "spring", "spring boot", "asp.net", "laravel", "ruby on rails",
    "rest", "rest api", "rest apis", "restful api", "restful apis", "graphql", "grpc", "microservices", "web sockets", "jwt", "oauth",
    # Databases & Storage
    "sql", "postgresql", "postgres", "mysql", "mongodb", "redis", "elasticsearch", "sqlite", "oracle", "sql server", "mssql", "cassandra", "dynamodb", "snowflake", "bigquery",
    # Cloud & DevOps
    "aws", "amazon web services", "azure", "microsoft azure", "gcp", "google cloud platform", "docker", "kubernetes", "k8s", "terraform", "jenkins",
    "ci/cd", "continuous integration", "continuous deployment", "github actions", "gitlab ci", "ansible", "helm", "linux", "bash", "shell scripting",
    # Data Science & AI / ML
    "machine learning", "deep learning", "nlp", "natural language processing", "computer vision", "tensorflow", "pytorch", "keras", "scikit-learn", "sklearn",
    "pandas", "numpy", "matplotlib", "seaborn", "scipy", "jupyter", "transformers", "bert", "sbert", "llm", "langchain", "opencv",
    # Big Data & Data Eng
    "apache spark", "spark", "apache kafka", "kafka", "airflow", "apache airflow", "etl", "data pipeline", "data warehousing", "hadoop",
    # QA & Testing
    "selenium", "cypress", "playwright", "junit", "testng", "postman", "jest", "pytest", "test automation", "unit testing",
    # Mobile
    "flutter", "react native", "android", "ios", "xcode", "android studio",
    # Security, Monitoring & UI/UX
    "figma", "sketch", "adobe xd", "wireframing", "prototyping", "prometheus", "grafana",
    "network security", "cybersecurity", "penetration testing", "siem", "splunk", "owasp",
    # Project & Tools
    "git", "github", "gitlab", "jira", "confluence", "agile", "scrum", "oop", "data structures", "algorithms"
]

EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")
PHONE_RE = re.compile(r"[\+]?[\d\s\-\(\)]{7,15}")
LINKEDIN_RE = re.compile(r"linkedin\.com/in/[\w\-]+", re.I)
GITHUB_RE = re.compile(r"github\.com/[\w\-]+", re.I)
YEARS_RE = re.compile(r"(\d+(?:\.\d+)?)\s*\+?\s*(?:years?\s+(?:of\s+)?(?:experience|work|professional))|(?:experience|worked)\s*:?\s*(\d+(?:\.\d+)?)\s*\+?\s*(?:years?|yrs?)", re.I)
EDU_RE = re.compile(r"\b(ph\.?d|doctorate|m\.?s\.?c?|master|b\.?s\.?c?|bachelor|b\.?tech|diploma|m\.?tech|mba)\b", re.I)
CERT_RE = re.compile(r"(?:certification|certificate|certified)[\s:]+([^\n]+)", re.I)
LANG_RE = re.compile(r"(?:language|fluent|proficient)[\s:]+([^\n]+)", re.I)

# ── Date patterns ───────────────────────────────────────────────
MONTH_MAP = {
    "jan": 1, "january": 1, "feb": 2, "february": 2, "mar": 3, "march": 3,
    "apr": 4, "april": 4, "may": 5, "jun": 6, "june": 6,
    "jul": 7, "july": 7, "aug": 8, "august": 8, "sep": 9, "september": 9,
    "oct": 10, "october": 10, "nov": 11, "november": 11, "dec": 12, "december": 12,
}
DATE_RANGE_RE = re.compile(
    r"((?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|"
    r"aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)"
    r"\s+\d{4})\s*[-–—to]+\s*"
    r"((?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|"
    r"aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)"
    r"\s+\d{4}|present|current|now)",
    re.I
)
# Also match "2022 - 2023" or "2022 - Present"
YEAR_RANGE_RE = re.compile(
    r"(\d{4})\s*[-–—to]+\s*(\d{4}|present|current|now)",
    re.I
)

# ── Project section patterns ────────────────────────────────────
PROJECT_SECTION_RE = re.compile(
    r"(?:academic\s+projects?|personal\s+projects?|projects?|capstone|thesis|"
    r"dissertation|coursework|portfolio|side\s+projects?|open[\s-]?source)"
    r"[\s:]*\n",
    re.I
)
ACADEMIC_KEYWORDS = {
    "thesis", "dissertation", "capstone", "coursework", "academic", "university",
    "college", "research", "paper", "study", "assignment", "lab", "seminar",
    "class", "school", "degree", "bachelor", "master", "phd",
}
PERSONAL_KEYWORDS = {
    "personal", "side project", "open source", "hobby", "freelance",
    "portfolio", "github", "contributed", "maintained",
}

# Lines that are education, NOT projects — skip these entirely
EDUCATION_LINE_RE = re.compile(
    r"(?:university|college|school|institute|academy|g\.?c\.?e|"
    r"\bo/?l\b|\ba/?l\b|"
    r"bachelor|master|phd|diploma|degree|passed\s+finalist|"
    r"education|certification|certificate|qualified|"
    r"\bb\.?sc\b|\bb\.?tech\b|\bm\.?sc\b|\bm\.?tech\b|\bph\.?d\b|\bb\.?e\b|\bm\.?ba\b)",
    re.I
)


def _parse_date(date_str: str) -> datetime | None:
    """Parse a date string like 'January 2023', 'Jan 2023', or just '2023'."""
    date_str = date_str.strip().lower()
    if date_str in ("present", "current", "now"):
        return datetime.now()
    parts = date_str.split()
    if len(parts) == 2:
        month_str, year_str = parts
        month = MONTH_MAP.get(month_str[:3])
        if month:
            try:
                return datetime(int(year_str), month, 1)
            except (ValueError, TypeError):
                pass
    elif len(parts) == 1:
        # Year only — assume January
        try:
            return datetime(int(parts[0]), 1, 1)
        except (ValueError, TypeError):
            pass
    return None


def _calc_duration_months(start: datetime, end: datetime) -> float:
    """Calculate duration in months between two dates."""
    months = (end.year - start.year) * 12 + (end.month - start.month)
    return max(months, 0.5)  # minimum 0.5 months (2 weeks)


_SECTION_RE = re.compile(
    r"^(?P<h>"
    r"academic\s+projects?|personal\s+projects?|side\s+projects?|open[\s-]?source\s+projects?|"
    r"research\s+projects?|projects?|capstone|thesis|dissertation|coursework|portfolio|"
    r"(?:work\s+|professional\s+|job\s+)?experience|employment|internship|"
    r"education|academic\s+background|qualifications?|"
    r"skills?|technical\s+skills?|core\s+competencies|soft\s+skills?|technolog(?:y|ies)\s+(?:and|&)\s+frameworks?|tech\s+stack|"
    r"certifications?|licenses?|languages?|"
    r"awards?|honou?rs?|achievements?|"
    r"interests?|hobbies?|references?|"
    r"activities?|leadership|clubs?(?:\s*&\s*leadership(?:\s*skills?)?)?|extra[- ]curricular|volunteering|"
    r"publications?|profile|summary|objective|about(\s+me)?"
    r")\s*:?\s*$",
    re.I
)


def _section_kind(header: str) -> str:
    """Normalize a section header to a canonical kind."""
    h = header.lower().strip().rstrip(":")
    if re.search(r"project|thesis|dissertation|capstone|coursework|portfolio", h):
        return "projects"
    if re.search(r"experience|employment|internship", h):
        return "experience"
    if re.search(r"education|background|qualification", h):
        return "education"
    if re.search(r"skill|competenc", h):
        return "skills"
    if re.search(r"certif|licens", h):
        return "certifications"
    if re.search(r"language", h):
        return "languages"
    if re.search(r"award|honou?r|achievement", h):
        return "awards"
    if re.search(r"interest|hobb", h):
        return "interests"
    if re.search(r"reference", h):
        return "references"
    if re.search(r"activit|leadership|extra|volunteer", h):
        return "activities"
    return "other"


def _segment_sections(text: str) -> list[tuple[str, list[str]]]:
    """Split resume text into (kind, lines) segments by section header."""
    segments: list[tuple[str, list[str]]] = []
    current_kind = "other"
    current_lines: list[str] = []
    for line in text.split("\n"):
        m = _SECTION_RE.match(line.strip())
        if m:
            if current_lines:
                segments.append((current_kind, current_lines))
            current_kind = _section_kind(m.group("h"))
            current_lines = [line]
        else:
            current_lines.append(line)
    if current_lines:
        segments.append((current_kind, current_lines))
    return segments


def _segment_text(segments: list[tuple[str, list[str]]], kind: str) -> str:
    """Join all lines belonging to a given section kind."""
    return "\n".join(line for k, lines in segments if k == kind for line in lines)


def _extract_projects_with_dates(text: str) -> tuple[list[dict], list[dict]]:
    """Extract academic and personal projects with date ranges.
    
    Returns:
        (academic_projects, personal_projects) — each is a list of dicts:
        [{"name": str, "duration_months": float, "dates": str}, ...]
    """
    lines = text.split("\n")
    academic_projects = []
    personal_projects = []
    seen = set()
    
    # Section header patterns to skip
    SECTION_HEADER_RE = re.compile(
        r"^(?:academic\s+|personal\s+|side\s+|open[\s-]?source\s+)?projects?\s*:?\s*$",
        re.I
    )
    SECTION_KIND_RE = re.compile(
        r"^(academic|personal|side|open[\s-]?source)\s+projects?\s*:?\s*$",
        re.I
    )
    # Non-project section headers that stop project extraction
    NON_PROJECT_HEADER_RE = re.compile(
        r"^(experience|education|skills?|certifications?|languages?|awards?|"
        r"interests?|references?|activities?|leadership)\s*:?\s*$",
        re.I
    )

    current_section = None  # "academic" | "personal" | None
    in_projects = False
    saw_header = False

    for i, line in enumerate(lines):
        line_stripped = line.strip()
        if not line_stripped:
            continue

        # Track which project section we're inside
        sk = SECTION_KIND_RE.match(line_stripped)
        if sk:
            sec = sk.group(1).lower()
            current_section = "academic" if "academic" in sec else "personal"
            in_projects = True
            saw_header = True
            continue
        if SECTION_HEADER_RE.match(line_stripped):
            in_projects = True
            saw_header = True
            continue
        if NON_PROJECT_HEADER_RE.match(line_stripped):
            in_projects = False
            current_section = None
            saw_header = True
            continue

        # Skip education entries (universities, schools, certificates)
        if EDUCATION_LINE_RE.search(line_stripped):
            continue

        # Skip club memberships and extra-curricular activities
        CLUB_RE = re.compile(r'\b(club|member|society|association|committee|volunteer|organization)\b', re.I)
        if CLUB_RE.search(line_stripped) and (DATE_RANGE_RE.search(line_stripped) or YEAR_RANGE_RE.search(line_stripped)):
            continue

        # Only treat as a project if we're inside a projects section
        # (fall back to whole-text when no headers were detected at all)
        if saw_header and not in_projects:
            continue

        # Look for date ranges on this line
        date_match = DATE_RANGE_RE.search(line_stripped)
        if not date_match:
            date_match = YEAR_RANGE_RE.search(line_stripped)
        
        # Check for single date (e.g., "March 2023") if no range found
        single_date = None
        if not date_match:
            single_date_match = re.search(
                r"((?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|"
                r"aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)"
                r"\s+\d{4})",
                line_stripped, re.I
            )
            if single_date_match:
                single_date = single_date_match.group(1)

        # Check for single year at end of line (e.g., "Project Name 2025")
        single_year = None
        if not date_match and not single_date:
            year_match = re.search(r'\b(19|20)\d{2}\s*$', line_stripped)
            if year_match:
                single_year = year_match.group(0).strip()

        if not date_match and not single_date and not single_year:
            continue

        # Extract the project name (the line with the date, or the line before)
        name_line = line_stripped
        # Remove date range from the name
        name_line = DATE_RANGE_RE.sub("", name_line).strip()
        name_line = YEAR_RANGE_RE.sub("", name_line).strip()
        # Remove single year at end of line
        name_line = re.sub(r'\b(19|20)\d{2}\s*$', '', name_line).strip()
        name_line = re.sub(r"^[-–—•\*\s]+", "", name_line).strip()
        name_line = re.sub(r"[-–—•\*\s]+$", "", name_line).strip()
        
        if not name_line or len(name_line) < 3:
            continue
        if name_line in seen:
            continue
        seen.add(name_line)

        # Parse dates — use the match from THIS line if available
        dm = date_match  # prefer the date on the same line
        
        if dm:
            groups = dm.groups()
            start_str = groups[0] if groups else ""
            end_str = groups[1] if len(groups) > 1 else ""
            start_date = _parse_date(start_str)
            end_date = _parse_date(end_str) if end_str else None
            if start_date and end_date:
                duration = _calc_duration_months(start_date, end_date)
                date_display = f"{start_str} - {end_str}"
            elif start_date:
                duration = 0
                date_display = start_str
            else:
                duration = 0
                date_display = ""
        elif single_date:
            # Single date found, no range
            start_date = _parse_date(single_date)
            duration = 0
            date_display = single_date if start_date else ""
        elif single_year:
            # Single year at end of line
            duration = 0
            date_display = single_year
        else:
            duration = 0
            date_display = ""

        project = {
            "name": name_line,
            "duration_months": duration,
            "dates": date_display,
        }

        # Classify as academic or personal (name keywords override section)
        name_lower = name_line.lower()
        is_academic = any(kw in name_lower for kw in ACADEMIC_KEYWORDS)
        is_personal = any(kw in name_lower for kw in PERSONAL_KEYWORDS)

        # Check the line right after the project name for "Personal Project" / "Coursework"
        if not is_academic and not is_personal:
            next_line = lines[i + 1].lower() if i + 1 < len(lines) else ""
            is_personal = "personal project" in next_line
            is_academic = any(k in next_line for k in ("coursework", "module", "year 1", "year 2", "year 3", "year 4"))

        if is_academic:
            academic_projects.append(project)
        elif is_personal:
            personal_projects.append(project)
        else:
            # Fall back to the section we're inside
            (academic_projects if current_section == "academic" else personal_projects).append(project)

    return academic_projects, personal_projects


def extract_entities(text: str) -> dict[str, Any]:
    """Extract structured entities from resume text with high accuracy."""
    entities: dict[str, Any] = {
        "name": "", "email": "", "phone": "", "address": "",
        "linkedin": "", "github": "", "skills": [], "education": "",
        "experience_years": 0, "projects": [],
        "academic_projects": [], "personal_projects": [],
        "project_experience_years": 0,
        "certifications": [], "languages": [], "tools": [], "frameworks": [],
    }

    phones = PHONE_RE.findall(text)
    if phones:
        entities["phone"] = phones[0].strip()

    emails = EMAIL_RE.findall(text)
    if emails:
        raw_email = emails[0]
        # Clean phone digits attached to start of email (e.g. 114inukajathmal11@gmail.com -> inukajathmal11@gmail.com)
        phone_digits = re.sub(r'\D', '', entities["phone"])
        if phone_digits:
            for d_len in range(6, 1, -1):
                suffix = phone_digits[-d_len:]
                if raw_email.startswith(suffix) and len(raw_email) > d_len and raw_email[d_len].isalpha():
                    raw_email = raw_email[d_len:]
                    break
        entities["email"] = raw_email

    # Candidate Name extraction
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    for line in lines[:10]:
        line_lower = line.lower()
        if any(sw in line_lower for sw in ['about', 'education', 'projects', 'experience', 'skills', 'summary', 'profile', 'curriculum', 'resume', 'contact', 'references']):
            continue
        if EMAIL_RE.search(line) or PHONE_RE.search(line) or 'http' in line_lower:
            continue
        words = line.split()
        if 1 <= len(words) <= 4 and all(re.match(r'^[A-Za-z\.]+$', w) for w in words):
            if line_lower in {'developer', 'engineer', 'manager', 'student', 'designer', 'analyst'}:
                continue
            entities["name"] = line.title()
            break

    if not entities["name"] and lines:
        first_line = lines[0]
        if not EMAIL_RE.search(first_line) and not PHONE_RE.search(first_line):
            entities["name"] = first_line[:100].title()

    linkedin = LINKEDIN_RE.findall(text)
    if linkedin:
        entities["linkedin"] = "https://" + linkedin[0]

    github = GITHUB_RE.findall(text)
    if github:
        entities["github"] = "https://" + github[0]

    segments = _segment_sections(text)

    text_lower = text.lower()
    text_normalized = re.sub(r"\s+", " ", text_lower)
    found_skills = []
    used_c1 = False
    if c1_extract_skills:
        try:
            c1_data = c1_extract_skills(text)
            detected = c1_data.get("detected_skills", [])
            found_skills = [s.title() for s in detected]
            if c1_data.get("detected_certs"):
                entities["certifications"] = list(dict.fromkeys(entities.get("certifications", []) + [c.title() for c in c1_data["detected_certs"]]))
            used_c1 = True
        except Exception:
            used_c1 = False

    if not used_c1:
        for skill in SKILLS_KEYWORDS:
            if re.search(r"(?:\b|_)" + re.escape(skill) + r"(?:\b|_)", text_normalized, re.I):
                found_skills.append(skill.title())

    entities["skills"] = list(dict.fromkeys(found_skills))

    # Education: extract complete degree specification line
    edu_text = _segment_text(segments, "education")
    edu_lines = []
    target_edu_texts = [edu_text, text] if edu_text else [text]
    for current_text in target_edu_texts:
        if not current_text:
            continue
        for line in current_text.split('\n'):
            clean_line = line.strip()
            if EDU_RE.search(clean_line):
                if not any(k in clean_line.lower() for k in ['passed finalist', 'ordinary level', 'advanced level', 'gce', 'school', 'scrum master']):
                    clean_line = re.sub(r'^[•\*\-\s]+', '', clean_line)
                    clean_line = re.sub(r'\|\s*\d{4}\s*[-–]\s*\d{4}', '', clean_line).strip()
                    if clean_line and len(clean_line) > 3 and not re.match(r'^(?:education|educational)\s*:?$', clean_line, re.I):
                        edu_lines.append(clean_line)
        if edu_lines:
            break

    if edu_lines:
        entities["education"] = " | ".join(edu_lines[:2])
    else:
        for current_text in target_edu_texts:
            edu_match = EDU_RE.search(current_text)
            if edu_match:
                entities["education"] = edu_match.group(0).title()
                break

    # Experience: Extract from work experience section and non-education lines
    has_exp_section = any(k == "experience" for k, _ in segments)
    exp_text = _segment_text(segments, "experience") if has_exp_section else ""

    explicit_years = []
    for pattern in [
        r'(\d+(?:\.\d+)?)\s*\+?\s*(?:years?|yrs?)\s+(?:of\s+)?(?:experience|work|industry|field|background)',
        r'(?:experience|worked|working)\s+(?:for\s+)?(\d+(?:\.\d+)?)\s*\+?\s*(?:years?|yrs?)',
        r'(\d+(?:\.\d+)?)\s*\+?\s*(?:years?|yrs?)\s+in\b',
        r'over\s+(\d+(?:\.\d+)?)\s*\+?\s*(?:years?|yrs?)',
    ]:
        for m in re.findall(pattern, exp_text or text, re.I):
            try:
                v = float(m)
                if 0.5 <= v <= 40.0:
                    explicit_years.append(v)
            except ValueError:
                pass

    # Extract date ranges from experience section (or lines not in education)
    target_exp_lines = exp_text.split('\n') if exp_text else [
        l for l in text.split('\n')
        if not EDUCATION_LINE_RE.search(l) and not any(k in l.lower() for k in ['university', 'college', 'bachelor', 'master', 'phd', 'degree', 'passed finalist', 'gce', 'advance level', 'ordinary level', 'diploma'])
    ]

    exp_intervals = []
    for line in target_exp_lines:
        line_clean = line.strip()
        if not line_clean:
            continue
        # Check date range or year range
        for dm in (DATE_RANGE_RE.finditer(line_clean), YEAR_RANGE_RE.finditer(line_clean)):
            for match in dm:
                g = match.groups()
                if len(g) >= 2:
                    st = _parse_date(g[0])
                    en = _parse_date(g[1])
                    if st and en:
                        st_yr = st.year + (st.month - 1) / 12.0
                        en_yr = en.year + (en.month - 1) / 12.0
                        if 1980 <= st_yr <= 2030 and st_yr <= en_yr and (en_yr - st_yr) <= 35:
                            exp_intervals.append((st_yr, en_yr))

    calc_work_years = 0.0
    if exp_intervals:
        exp_intervals.sort(key=lambda x: x[0])
        merged = []
        for st, en in exp_intervals:
            if not merged:
                merged.append([st, en])
            else:
                prev = merged[-1]
                if st <= prev[1]:
                    prev[1] = max(prev[1], en)
                else:
                    merged.append([st, en])
        total_span = sum(en - st for st, en in merged)
        if total_span > 0:
            calc_work_years = round(total_span, 1)

    final_exp = calc_work_years
    if explicit_years:
        final_exp = max(max(explicit_years), calc_work_years)

    entities["experience_years"] = min(final_exp, 40.0)

    # Projects: only look inside the PROJECTS section
    proj_text = _segment_text(segments, "projects") or text
    academic_projects, personal_projects = _extract_projects_with_dates(proj_text)
    entities["academic_projects"] = academic_projects
    entities["personal_projects"] = personal_projects

    # Also keep flat project names for backward compat
    all_project_names = [p["name"] for p in academic_projects + personal_projects]
    entities["projects"] = all_project_names[:10]

    # Calculate total project experience in years
    total_months = sum(p["duration_months"] for p in academic_projects + personal_projects)
    entities["project_experience_years"] = round(total_months / 12, 1) if total_months > 0 else 0

    # Add project experience to experience_years if no work experience found
    if entities["experience_years"] == 0 and entities["project_experience_years"] > 0:
        entities["experience_years"] = entities["project_experience_years"]

    certs = CERT_RE.findall(text)
    entities["certifications"] = [c.strip() for c in certs[:5]]

    langs = LANG_RE.findall(text)
    entities["languages"] = [l.strip() for l in langs[:5]]

    return entities
